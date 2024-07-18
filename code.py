import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from datetime import datetime, timedelta

# Load data
df = pd.read_csv('finaldata.csv')

# Initialize session state for login and data
if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False
if 'last_entered_data' not in st.session_state:
    st.session_state['last_entered_data'] = None
if 'last_submit_time' not in st.session_state:
    st.session_state['last_submit_time'] = None
if 'input_data' not in st.session_state:
    st.session_state['input_data'] = {}

# Function to check credentials
def check_credentials(username, password):
    return username == "username" and password == "password"

# Login and Signup
if not st.session_state['authenticated']:
    st.title('Hospital Data Analytics Login')
    
    # Login form
    with st.form("login_form"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submit = st.form_submit_button("Login")

        if submit:
            if check_credentials(username, password):
                st.session_state['authenticated'] = True
                st.experimental_rerun()
            else:
                st.error("Invalid username or password")
else:
    st.title('Hospital Data Analytics')
    # Create tabs
    tab1, tab2, tab3 = st.tabs(["Data Entry", "Data Analytics", "Generate Report"])

    # Tab 1: Data Entry
    with tab1:
        st.header('Data Entry')

      

        # Create input fields for each column
        for column in df.columns:
            if 'Date' in column or column in ['In date', 'Out date']:
                st.session_state['input_data'][column] = st.date_input(f'Select {column}', value=st.session_state['input_data'].get(column, None))
            else:
                st.session_state['input_data'][column] = st.text_input(f'Enter {column}', value=st.session_state['input_data'].get(column, ''))

        if st.button('Submit Data'):
            if all(st.session_state['input_data'].values()):
                # Convert dates to string format
                for column in df.columns:
                    if 'Date' in column or column in ['In date', 'Out date']:
                        st.session_state['input_data'][column] = st.session_state['input_data'][column].strftime('%Y-%m-%d')
                
                new_data = {col: [st.session_state['input_data'][col]] for col in df.columns}
                new_df = pd.DataFrame(new_data)
                df = pd.concat([df, new_df], ignore_index=True)
                df.to_csv('C:\\Users\\nanis\\OneDrive\\Desktop\\hda\\finaldata.csv', index=False)
                st.session_state['last_entered_data'] = new_df
                st.session_state['last_submit_time'] = datetime.now()
                st.success('Data submitted successfully')

                # Clear input fields after submission
                st.session_state['input_data'].clear()
                st.experimental_rerun()
            else:
                st.error('Please fill in all fields before submitting')

        if st.session_state['last_entered_data'] is not None:
            time_diff = datetime.now() - st.session_state['last_submit_time']
            if time_diff < timedelta(minutes=2):
                st.write('**Last Entered Data Overview:**')
                st.write(st.session_state['last_entered_data'])
                st.write('---')
            else:
                st.session_state['last_entered_data'] = None

    # Tab 2: Data Analytics
    with tab2:
        st.header('Data Analytics')
        severity_counts = df.groupby(['PatientRegion', 'Severity']).size().reset_index(name='Count')

        charts = [
            {
                'title': 'Department-wise Patient Distribution',
                'chart': px.pie(df, names='Department'),
                'conclusion': 'Most patients are admitted to the Emergency and Cardiology departments.'
            },
            {
                'title': 'Age Distribution',
                'chart': px.histogram(df, x='PatientAge'),
                'conclusion': 'The age distribution of patients shows that the majority are between 30 and 50 years old.'
            },
            {
                'title': 'Gender Distribution',
                'chart': px.pie(df, names='PatientGender'),
                'conclusion': 'The gender distribution indicates a relatively equal representation of male and female patients.'
            },
            {
                'title': 'Region-wise Distribution',
                'chart': px.bar(df, x='PatientRegion', y='Final bill'),
                'conclusion': 'Patients from Region A have the highest total final bill amount, indicating a larger patient base or higher treatment costs in that region.'
            },
             {
            'title': 'Severity of Cases from Different Regions',
            'chart': px.bar(severity_counts, x='PatientRegion', y='Count', color='Severity', barmode='stack',
                        labels={'Count': 'Number of Cases', 'PatientRegion': 'Region', 'Severity': 'Severity Level'},
                        title='Severity of Cases from Different Regions'),
            'conclusion': 'The severity of cases varies significantly across different regions, with some regions experiencing more severe cases than others.'
            },
            {
                'title': 'Insurance Coverage',
                'chart': px.scatter(df, x='Insurance', y='Final bill', color='Severity'),
                'conclusion': 'There is a positive correlation between the amount covered by insurance and the final bill, particularly for severe cases.'
            },
            {
                'title': 'Doctor Consulted Distribution',
                'chart': px.histogram(df, x='Doctor consulted'),
                'conclusion': 'The distribution of patients among doctors shows some doctors handle significantly more cases than others.'
            },
            
            {
                'title': 'Final Bill by Age Group',
                'chart': px.box(df, x='PatientAge', y='Final bill', color='PatientGender'),
                'conclusion': 'Older patients generally have higher final bills, especially among males.'
            }
        ]

        if 'chart_index' not in st.session_state:
            st.session_state.chart_index = 0

        def show_chart():
            st.subheader(charts[st.session_state.chart_index]['title'])
            st.plotly_chart(charts[st.session_state.chart_index]['chart'])

        show_chart()

        if st.session_state.chart_index > 0:
            if st.button('Previous Chart'):
                st.session_state.chart_index -= 1

        if st.session_state.chart_index < len(charts) - 1:
            if st.button('Next Chart'):
                st.session_state.chart_index += 1
        else:
            st.info('No more charts to display.')

    # Tab 3: Generate Report
    with tab3:
        st.header('Generate Report')
        
        # Date input
        start_date = st.date_input('Start Date')
        end_date = st.date_input('End Date')

        def generate_report(start_date, end_date):
            try:
                # Ensure start_date and end_date are strings
                start_date = start_date.strftime('%Y-%m-%d')
                end_date = end_date.strftime('%Y-%m-%d')

                filtered_df = df[(df['In date'] >= start_date) & (df['Out date'] <= end_date)]
                
                # Calculations for the report
                total_patients = filtered_df.shape[0]
                avg_age = filtered_df['PatientAge'].mean()
                gender_distribution = filtered_df['PatientGender'].value_counts()
                total_final_bill = filtered_df['Final bill'].sum()
                severity = filtered_df['Severity'].value_counts()
                most_severe = severity.idxmax() if not severity.empty else 'N/A'
                most_severe_count = severity.max() if not severity.empty else 0
                patient_region = filtered_df['PatientRegion'].value_counts()
                most_common_region = patient_region.idxmax() if not patient_region.empty else 'N/A'
                most_common_region_count = patient_region.max() if not patient_region.empty else 0

                # Generate the report document
                doc = Document()
                doc.add_heading('HOSPITAL ANALYTICS REPORT', 0)
                doc.add_heading(f'Date Range: {start_date} to {end_date}', level=1)

                doc.add_heading('CONCLUSIONS:', level=1)
                doc.add_paragraph(f'Total patients: {total_patients}')
                doc.add_paragraph(f'Average age: {avg_age:.2f}')
                doc.add_paragraph(f'Gender distribution: {gender_distribution.to_dict()}')
                doc.add_paragraph(f'Total final bill: {total_final_bill:.2f}')
                doc.add_paragraph(f'Most common severity: {most_severe} with {most_severe_count} cases')

                # Add conclusion statements
                doc.add_heading('Detailed Analysis and Conclusions', level=1)
                doc.add_paragraph(f'The hospital admitted a total of {total_patients} patients between {start_date} and {end_date}.')
                doc.add_paragraph(f'The average age of the patients was {avg_age:.2f} years.')
                doc.add_paragraph(f'The gender distribution among the patients was as follows: {gender_distribution.to_dict()}.')
                doc.add_paragraph(f'The most common severity level observed was "{most_severe}" with {most_severe_count} cases.')

                # Adding more detailed conclusions
                if most_common_region_count > 5:  # Example threshold
                    doc.add_paragraph(f"Consider opening a new branch in {most_common_region} due to high patient inflow.")
                if most_severe_count > 2:  # Example threshold
                    doc.add_paragraph(f"Additional resources may be needed to handle the high number of {most_severe} cases.")

                # Additional conclusions
                doc.add_paragraph("The analysis indicates a significant number of patients are from the region "
                                  f"{most_common_region}, suggesting the need for targeted health initiatives in this area.")
                doc.add_paragraph(f"Patients with {most_severe} severity require more attention and resources. "
                                  "It is recommended to allocate specialized medical staff to manage these cases effectively.")
                doc.add_paragraph("Gender distribution shows a balanced ratio, indicating no gender-specific health trends "
                                  "within the analyzed period.")
                doc.add_paragraph("The hospital has generated a total revenue of ${:.2f} from the patients treated during the "
                                  "selected period.".format(total_final_bill))

                # Recommendations section
                doc.add_heading('RECOMMENDATIONS', level=1)
                doc.add_paragraph("1. Increase staffing levels during peak periods to handle the high patient inflow efficiently.")
                doc.add_paragraph(f"2. Consider expanding services or infrastructure in {most_common_region} to cater to the "
                                  "growing patient base.")
                doc.add_paragraph("3. Implement specialized training programs for medical staff to better handle severe cases.")
                doc.add_paragraph("4. Review and optimize billing processes to ensure accurate and timely revenue collection.")

                doc_file = 'monthly_Report.docx'
                doc.save(doc_file)
                st.success(f'Report generated and saved as {doc_file}')
            except Exception as e:
                st.error(f"Error generating report: {e}")

        if st.button('Generate Report'):
            if start_date and end_date:
                generate_report(start_date, end_date)
            else:
                st.error("Please select both start and end dates")
